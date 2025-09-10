"""
新しいRPAシステム

このプログラムは、ユーザ名、型番、製造番号、受注番号を入力し、
ExcelファイルとWordファイルの自動処理を行うRPAシステムです。

【入力項目】
- ユーザ名: 日本語の文字列
- 型番: 100,200,201,350,351の5種類+"-"+4桁数字+"."+6桁数字
- 製造番号: J000から始まり、4桁の数字、最後に0n00（nは1-9）
- 受注番号: O,N,Tのいずれか+4桁数字

【対応機種】
- 100, 200, 201: チェーンデータ31.75、データ1: 20
- 350, 351: チェーンデータ50.8、データ1: 14

【機能】
- 入力データの検証
- Excelファイル（check1.xlsx）への自動書き込み
- Wordファイル（check2.docx）の自動処理
- コンソール出力とGUI表示
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import re
import openpyxl
from openpyxl import load_workbook
import os
from datetime import datetime
from docx import Document


class NewRPASystem:
    """
    新しいRPAシステムのメインクラス

    ユーザ名、型番、製造番号、受注番号を入力し、
    ExcelファイルとWordファイルの自動処理を行うGUIアプリケーションを提供します。
    """
    
    # 定数定義
    SUPPORTED_MODELS = ["100", "200", "201", "350", "351"]
    MODEL_PATTERN = r"^(100|200|201|350|351)-\d{4}\.\d{6}$"
    MANUFACTURING_PATTERN = r"^J000\d{4}0[1-9]00$"
    ORDER_PATTERN = r"^[ONT]\d{4}$"
    
    # デフォルトファイル名
    DEFAULT_EXCEL_FILE = "check1.xlsx"
    DEFAULT_WORD_FILE = "check2.docx"
    
    # 置換対象キー文字列
    REPLACEMENT_KEY = "検査対象情報"

    def __init__(self):
        """NewRPASystemの初期化"""
        self.setup_gui()

    # ==================== GUI設定メソッド ====================

    def setup_gui(self):
        """GUIの設定"""
        self.root = tk.Tk()
        self.root.title("ファイル自動処理システム")
        self.root.geometry("700x600")
        self.root.configure(bg="#f0f0f0")

        # ウィンドウを中央に配置
        self.center_window()

        # スタイルの設定
        self.setup_styles()

        # メインフレーム
        main_frame = ttk.Frame(self.root, padding="30")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # ヘッダーセクション
        self.create_header(main_frame)

        # 入力フィールド
        self.create_input_fields(main_frame)

        # ボタン
        self.create_buttons(main_frame)

        # 結果表示
        self.create_result_display(main_frame)

        # グリッドの重み設定
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(4, weight=1)
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

    def center_window(self):
        """ウィンドウを画面中央に配置"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f"{width}x{height}+{x}+{y}")

    def setup_styles(self):
        """スタイルの設定"""
        style = ttk.Style()

        # テーマの設定
        style.theme_use("clam")

        # カスタムスタイルの定義
        style.configure("Title.TLabel", font=("Arial", 20, "bold"), foreground="#2c3e50", background="#f0f0f0")

        style.configure("Header.TLabel", font=("Arial", 12, "bold"), foreground="#34495e", background="#f0f0f0")

        style.configure("Field.TLabel", font=("Arial", 11), foreground="#2c3e50", background="#f0f0f0")

        style.configure("Custom.TEntry", fieldbackground="white", borderwidth=2, relief="solid", font=("Arial", 11))

        style.configure("Custom.TCombobox", fieldbackground="white", borderwidth=2, relief="solid", font=("Arial", 11))

        style.configure(
            "Primary.TButton",
            font=("Arial", 12, "bold"),
            foreground="white",
            background="#3498db",
            borderwidth=0,
            focuscolor="none",
        )

        style.configure(
            "Secondary.TButton",
            font=("Arial", 11),
            foreground="#2c3e50",
            background="#ecf0f1",
            borderwidth=1,
            focuscolor="none",
        )

        style.configure("Info.TLabelframe", background="#f0f0f0", borderwidth=1, relief="solid")

        style.configure(
            "Info.TLabelframe.Label", font=("Arial", 11, "bold"), foreground="#2c3e50", background="#f0f0f0"
        )

    def create_header(self, parent):
        """ヘッダーセクションの作成"""
        header_frame = ttk.Frame(parent)
        header_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 30))

        # タイトル
        title_label = ttk.Label(header_frame, text="新しいRPAシステム", style="Title.TLabel")
        title_label.grid(row=0, column=0, pady=(0, 10))

        # サブタイトル
        subtitle_label = ttk.Label(header_frame, text="データ入力・検証システム", style="Header.TLabel")
        subtitle_label.grid(row=1, column=0)

    def create_input_fields(self, parent):
        """入力フィールドの作成"""
        # 入力フィールドフレーム
        input_frame = ttk.LabelFrame(parent, text="入力情報", style="Info.TLabelframe", padding="20")
        input_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 20))

        # ユーザ名
        ttk.Label(input_frame, text="ユーザ名:", style="Field.TLabel").grid(row=0, column=0, sticky=tk.W, pady=(0, 5))
        self.username_var = tk.StringVar()
        username_entry = ttk.Entry(input_frame, textvariable=self.username_var, style="Custom.TEntry", width=35)
        username_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=(0, 5))

        # 型番
        ttk.Label(input_frame, text="型番:", style="Field.TLabel").grid(row=1, column=0, sticky=tk.W, pady=(0, 5))
        self.model_var = tk.StringVar()
        model_entry = ttk.Entry(input_frame, textvariable=self.model_var, style="Custom.TEntry", width=35)
        model_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=(0, 5))

        # 製造番号
        ttk.Label(input_frame, text="製造番号:", style="Field.TLabel").grid(row=2, column=0, sticky=tk.W, pady=(0, 5))
        self.manufacturing_var = tk.StringVar()
        manufacturing_entry = ttk.Entry(
            input_frame, textvariable=self.manufacturing_var, style="Custom.TEntry", width=35
        )
        manufacturing_entry.grid(row=2, column=1, sticky=(tk.W, tk.E), pady=(0, 5))

        # 受注番号
        ttk.Label(input_frame, text="受注番号:", style="Field.TLabel").grid(row=3, column=0, sticky=tk.W, pady=(0, 5))
        self.order_var = tk.StringVar()
        order_entry = ttk.Entry(input_frame, textvariable=self.order_var, style="Custom.TEntry", width=35)
        order_entry.grid(row=3, column=1, sticky=(tk.W, tk.E), pady=(0, 5))

        # グリッドの重み設定
        input_frame.columnconfigure(1, weight=1)

        # 入力例の表示
        self.create_input_examples(parent)

    def create_input_examples(self, parent):
        """入力例の表示"""
        examples_frame = ttk.LabelFrame(parent, text="入力例・ルール説明", style="Info.TLabelframe", padding="15")
        examples_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 20))

        examples_text = """【入力例】
ユーザ名: マキシンコー
型番: 201-2312.003000
製造番号: J00012340100
受注番号: O1234

"""

        examples_label = ttk.Label(
            examples_frame, text=examples_text, font=("Consolas", 9), justify=tk.LEFT, foreground="#2c3e50"
        )
        examples_label.grid(row=0, column=0, sticky=tk.W)

    def create_buttons(self, parent):
        """ボタンの作成"""
        button_frame = ttk.Frame(parent)
        button_frame.grid(row=3, column=0, columnspan=2, pady=(0, 20))

        # 実行ボタン
        execute_button = ttk.Button(button_frame, text="🚀 実行", command=self.execute_rpa, style="Primary.TButton")
        execute_button.grid(row=0, column=0, padx=(0, 5))

        # クリアボタン
        clear_button = ttk.Button(button_frame, text="🗑️ クリア", command=self.clear_inputs, style="Secondary.TButton")
        clear_button.grid(row=0, column=1, padx=(5, 5))

        # Excel書き込みボタン
        excel_button = ttk.Button(
            button_frame, text="📊 Excel書き込み", command=self.write_to_excel_direct, style="Primary.TButton"
        )
        excel_button.grid(row=0, column=2, padx=(5, 5))

        # Word処理ボタン
        word_button = ttk.Button(
            button_frame, text="📝 Word処理", command=self.process_word_direct, style="Primary.TButton"
        )
        word_button.grid(row=0, column=3, padx=(5, 0))

    def create_result_display(self, parent):
        """結果表示エリアの作成"""
        result_frame = ttk.LabelFrame(parent, text="📊 実行結果", style="Info.TLabelframe", padding="15")
        result_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))

        self.result_text = tk.Text(
            result_frame,
            height=10,
            width=70,
            font=("Consolas", 10),
            wrap=tk.WORD,
            bg="#ffffff",
            fg="#2c3e50",
            relief="solid",
            borderwidth=1,
            padx=10,
            pady=10,
        )
        self.result_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # スクロールバー
        scrollbar = ttk.Scrollbar(result_frame, orient=tk.VERTICAL, command=self.result_text.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.result_text.configure(yscrollcommand=scrollbar.set)

        result_frame.columnconfigure(0, weight=1)
        result_frame.rowconfigure(0, weight=1)

    # ==================== イベントハンドラメソッド ====================

    def execute_rpa(self):
        """RPA実行"""
        try:
            # 入力値の取得
            username = self.username_var.get().strip()
            model = self.model_var.get().strip()
            manufacturing = self.manufacturing_var.get().strip()
            order = self.order_var.get().strip()

            # 入力値の検証
            validation_result = self.validate_inputs(username, model, manufacturing, order)
            if not validation_result["valid"]:
                messagebox.showerror("入力エラー", validation_result["message"])
                return

            # コンソール出力
            self.output_to_console(username, model, manufacturing, order)

            # 結果表示エリアに出力
            self.display_result(username, model, manufacturing, order)

            messagebox.showinfo("実行完了", "RPAが正常に実行されました。")

        except Exception as e:
            messagebox.showerror("エラー", f"実行中にエラーが発生しました: {str(e)}")

    def clear_inputs(self):
        """入力フィールドのクリア"""
        self.username_var.set("")
        self.model_var.set("")
        self.manufacturing_var.set("")
        self.order_var.set("")
        self.result_text.delete(1.0, tk.END)

    # ==================== 検証メソッド ====================

    def validate_username(self, username):
        """
        ユーザ名の検証

        Args:
            username (str): ユーザ名

        Returns:
            dict: 検証結果 {"valid": bool, "message": str}
        """
        if not username:
            return {"valid": False, "message": "ユーザ名を入力してください。"}

        if not username.strip():
            return {"valid": False, "message": "ユーザ名は空白のみでは入力できません。"}

        # 日本語文字のチェック（ひらがな、カタカナ、漢字、英数字、記号）
        import re

        if not re.search(r"[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FAF\uFF00-\uFFEF]", username):
            return {"valid": False, "message": "ユーザ名には日本語文字を含む必要があります。"}

        return {"valid": True, "message": ""}

    def validate_model(self, model):
        """
        型番の検証

        Args:
            model (str): 型番

        Returns:
            dict: 検証結果 {"valid": bool, "message": str}
        """
        if not model:
            return {"valid": False, "message": "型番を入力してください。"}

        # 型番の形式チェック
        if not re.match(self.MODEL_PATTERN, model):
            return {
                "valid": False,
                "message": "型番の形式が正しくありません。\n形式: 100,200,201,350,351のいずれか-4桁数字.6桁数字\n例: 100-2312.003000",
            }

        return {"valid": True, "message": ""}

    def validate_manufacturing_number(self, manufacturing):
        """
        製造番号の検証

        Args:
            manufacturing (str): 製造番号

        Returns:
            dict: 検証結果 {"valid": bool, "message": str}
        """
        if not manufacturing:
            return {"valid": False, "message": "製造番号を入力してください。"}

        # 製造番号の形式チェック
        if not re.match(self.MANUFACTURING_PATTERN, manufacturing):
            return {
                "valid": False,
                "message": "製造番号の形式が正しくありません。\n形式: J000+4桁数字+0+1-9+00\n例: J00023150100",
            }

        return {"valid": True, "message": ""}

    def validate_order_number(self, order):
        """
        受注番号の検証

        Args:
            order (str): 受注番号

        Returns:
            dict: 検証結果 {"valid": bool, "message": str}
        """
        if not order:
            return {"valid": False, "message": "受注番号を入力してください。"}

        # 受注番号の形式チェック
        if not re.match(self.ORDER_PATTERN, order):
            return {"valid": False, "message": "受注番号の形式が正しくありません。\n形式: O/N/T+4桁数字\n例: O2315"}

        return {"valid": True, "message": ""}

    def validate_manufacturing_order_consistency(self, manufacturing, order):
        """
        製造番号と受注番号の整合性検証

        Args:
            manufacturing (str): 製造番号
            order (str): 受注番号

        Returns:
            dict: 検証結果 {"valid": bool, "message": str}
        """
        # 製造番号の5-8文字目（J000の後の4桁数字）と受注番号の1-4文字目（O/N/T+3桁数字）が一致する必要がある
        manufacturing_middle = manufacturing[4:8]  # J000の後の4桁数字（例：2315）
        order_prefix = order[1:5]  # 受注番号の1-4文字目（O/N/T+3桁数字）（例：2315）

        if manufacturing_middle != order_prefix:
            return {
                "valid": False,
                "message": f"製造番号と受注番号が一致しません。\n製造番号の5-8文字目: {manufacturing_middle}\n受注番号の1-4文字目: {order_prefix}\nこれらは同じである必要があります。",
            }

        return {"valid": True, "message": ""}

    def validate_inputs(self, username, model, manufacturing, order):
        """
        入力値の検証（包括的）

        Args:
            username (str): ユーザ名
            model (str): 型番
            manufacturing (str): 製造番号
            order (str): 受注番号

        Returns:
            dict: 検証結果 {"valid": bool, "message": str}
        """
        # 各フィールドの個別検証
        validations = [
            ("ユーザ名", self.validate_username(username)),
            ("型番", self.validate_model(model)),
            ("製造番号", self.validate_manufacturing_number(manufacturing)),
            ("受注番号", self.validate_order_number(order)),
        ]

        # 個別検証でエラーがあれば最初のエラーを返す
        for field_name, result in validations:
            if not result["valid"]:
                return {"valid": False, "message": f"{field_name}: {result['message']}"}

        # 製造番号と受注番号の整合性検証
        consistency_result = self.validate_manufacturing_order_consistency(manufacturing, order)
        if not consistency_result["valid"]:
            return consistency_result

        return {"valid": True, "message": ""}

    # ==================== 出力メソッド ====================

    def output_to_console(self, username, model, manufacturing, order):
        """
        コンソールに出力

        Args:
            username (str): ユーザ名
            model (str): 型番
            manufacturing (str): 製造番号
            order (str): 受注番号
        """
        print("=" * 50)
        print("新しいRPAシステム - 実行結果")
        print("=" * 50)
        print(f"ユーザ名: {username}")
        print(f"型番: {model}")
        print(f"製造番号: {manufacturing}")
        print(f"受注番号: {order}")
        print("=" * 50)

    def display_result(self, username, model, manufacturing, order):
        """
        結果表示エリアに出力

        Args:
            username (str): ユーザ名
            model (str): 型番
            manufacturing (str): 製造番号
            order (str): 受注番号
        """
        result = f"""✅ 実行完了

📋 入力データ:
┌─────────────────────────────────────────┐
│ ユーザ名: {username:<25} │
│ 型番: {model:<27} │
│ 製造番号: {manufacturing:<23} │
│ 受注番号: {order:<25} │
└─────────────────────────────────────────┘

⏰ 実行時刻: {self.get_current_time()}

🎉 すべての検証が正常に完了しました！"""

        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(1.0, result)

    # ==================== ヘルパーメソッド ====================
    
    def get_current_time(self):
        """現在時刻の取得"""
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    def select_excel_file(self):
        """Excelファイルを選択"""
        file_path = filedialog.askopenfilename(
            title="Excelファイルを選択", filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        return file_path if file_path else None
    
    def select_word_file(self):
        """Wordファイルを選択"""
        file_path = filedialog.askopenfilename(
            title="Wordファイルを選択", filetypes=[("Word files", "*.docx *.doc"), ("All files", "*.*")]
        )
        return file_path if file_path else None

    # ==================== Excel操作メソッド ====================

    def write_to_excel(self, username, model, manufacturing, order, file_path=None):
        """Excelファイルにデータを書き込み"""
        try:
            # デフォルトファイルを使用
            if not file_path:
                file_path = self.DEFAULT_EXCEL_FILE
                if not os.path.exists(file_path):
                    messagebox.showerror("エラー", f"ファイルが見つかりません: {file_path}\n既存のExcelファイルを配置してください。")
                    return False

            # Excelファイルを開く
            wb = load_workbook(file_path)

            # 組立チェック表の書き込み
            if "組立チェック表" in wb.sheetnames:
                ws1 = wb["組立チェック表"]
                ws1["B4"] = f"ユーザー名：{username}"
                ws1["B5"] = f"機種-型番：{model}"
                ws1["F4"] = f"受注番号：{order}"
                ws1["F5"] = f"製造番号：{manufacturing}"

            # フレームテスト検査表の書き込み
            if "フレームテスト検査表" in wb.sheetnames:
                ws2 = wb["フレームテスト検査表"]
                ws2["B3"] = f"ユーザー名：{username}"
                ws2["B4"] = f"機種-型番：{model}"
                ws2["F2"] = f"受注番号：{order}"
                ws2["F3"] = f"製造番号：{manufacturing}"

            # フレーム組立検査表の書き込み
            if "フレーム組立検査表" in wb.sheetnames:
                ws3 = wb["フレーム組立検査表"]
                ws3["B3"] = f"ユーザー名：{username}"
                ws3["B4"] = f"機種-型番：{model}"
                ws3["F2"] = f"受注番号：{order}"
                ws3["F3"] = f"製造番号：{manufacturing}"

            # 新しいファイル名を生成（タイムスタンプ付き）
            base_name = os.path.splitext(file_path)[0]
            extension = os.path.splitext(file_path)[1]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            new_file_path = f"{base_name}_processed_{timestamp}{extension}"
            
            # 新しいファイルとして保存
            wb.save(new_file_path)

            messagebox.showinfo("成功", f"Excelファイルにデータを書き込みました:\n元ファイル: {os.path.basename(file_path)}\n新ファイル: {os.path.basename(new_file_path)}")
            return True, new_file_path

        except Exception as e:
            messagebox.showerror("エラー", f"Excelファイルの書き込みに失敗しました:\n{str(e)}")
            return False, None

    def write_to_excel_direct(self):
        """入力データを直接Excelに書き込み"""
        try:
            # 入力値の取得
            username = self.username_var.get().strip()
            model = self.model_var.get().strip()
            manufacturing = self.manufacturing_var.get().strip()
            order = self.order_var.get().strip()

            # 入力値の検証
            validation_result = self.validate_inputs(username, model, manufacturing, order)
            if not validation_result["valid"]:
                messagebox.showerror("入力エラー", validation_result["message"])
                return

            # Excelファイルに書き込み
            success, new_file_path = self.write_to_excel(username, model, manufacturing, order)

            if success:
                # 結果表示エリアに更新
                self.display_result(username, model, manufacturing, order)

                # 追加のExcel書き込み完了メッセージ
                excel_result = f"""\n\n📊 Excel書き込み完了:
✅ 元ファイル: check1.xlsx
✅ 新ファイル: {os.path.basename(new_file_path)}
✅ 組立チェック表: B4,B5,F4,F5セルにデータを書き込み
✅ フレームテスト検査表: B3,B4,F2,F3セルにデータを書き込み  
✅ フレーム組立検査表: B3,B4,F2,F3セルにデータを書き込み"""

                self.result_text.insert(tk.END, excel_result)

        except Exception as e:
            messagebox.showerror("エラー", f"Excel書き込み中にエラーが発生しました: {str(e)}")

    # ==================== Word操作メソッド ====================

    def replace_text_in_word(self, file_path, search_text, replace_text):
        """Wordファイル内のテキストを置換"""
        try:
            # Wordファイルを開く
            doc = Document(file_path)

            # 置換回数をカウント
            replacement_count = 0

            # 段落内のテキストを置換
            for paragraph in doc.paragraphs:
                if search_text in paragraph.text:
                    # 段落内のすべてのrunをチェック
                    for run in paragraph.runs:
                        if search_text in run.text:
                            run.text = run.text.replace(search_text, replace_text)
                            replacement_count += 1

            # テーブル内のテキストも置換
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if search_text in run.text:
                                    run.text = run.text.replace(search_text, replace_text)
                                    replacement_count += 1

            # テキストボックスや図形内のテキストも置換
            # 1. インライン図形
            for shape in doc.inline_shapes:
                if hasattr(shape, 'text_frame') and shape.text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        for run in paragraph.runs:
                            if search_text in run.text:
                                run.text = run.text.replace(search_text, replace_text)
                                replacement_count += 1
            
            # 2. フローティング図形（テキストボックス、図形など）
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # 図形内のテキストを検索
                    if hasattr(run, '_element') and run._element.tag.endswith('drawing'):
                        # 図形内のテキストを処理
                        for text_elem in run._element.iter():
                            if text_elem.tag.endswith('t'):  # テキスト要素
                                if search_text in text_elem.text:
                                    text_elem.text = text_elem.text.replace(search_text, replace_text)
                                    replacement_count += 1
            
            # 3. XMLレベルでの包括的な検索（図形内のテキスト）
            try:
                # 文書全体のXMLを検索
                for elem in doc._element.iter():
                    if elem.tag.endswith('t'):  # テキスト要素
                        if elem.text and search_text in elem.text:
                            elem.text = elem.text.replace(search_text, replace_text)
                            replacement_count += 1
            except Exception as xml_error:
                print(f"XML検索エラー（無視）: {xml_error}")

            # ヘッダーとフッター内のテキストも置換
            for section in doc.sections:
                # ヘッダー
                if section.header:
                    for paragraph in section.header.paragraphs:
                        for run in paragraph.runs:
                            if search_text in run.text:
                                run.text = run.text.replace(search_text, replace_text)
                                replacement_count += 1
                
                # フッター
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        for run in paragraph.runs:
                            if search_text in run.text:
                                run.text = run.text.replace(search_text, replace_text)
                                replacement_count += 1

            # 新しいファイル名を生成（タイムスタンプ付き）
            base_name = os.path.splitext(file_path)[0]
            extension = os.path.splitext(file_path)[1]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            new_file_path = f"{base_name}_processed_{timestamp}{extension}"
            
            # 新しいファイルとして保存
            doc.save(new_file_path)

            return replacement_count, new_file_path

        except Exception as e:
            print(f"Wordファイル処理エラー: {str(e)}")
            raise Exception(f"Wordファイルの処理中にエラーが発生しました: {str(e)}")



    def process_word_file(self, username, model, manufacturing, order, file_path=None):
        """Wordファイルを処理してキー文字列を置換"""
        try:
            # デフォルトファイルを使用
            if not file_path:
                file_path = self.DEFAULT_WORD_FILE
                if not os.path.exists(file_path):
                    messagebox.showerror("エラー", f"ファイルが見つかりません: {file_path}")
                    return False

            # 置換用のテキストを生成
            replacement_text = f"{order}/{manufacturing}"

            # 置換対象のキー文字列
            key_strings = [self.REPLACEMENT_KEY]

            total_replacements = 0

            # キー文字列を置換
            new_file_path = None
            for key_string in key_strings:
                count, new_file_path = self.replace_text_in_word(file_path, key_string, replacement_text)
                total_replacements += count

            if total_replacements > 0:
                messagebox.showinfo(
                    "成功",
                    f"Wordファイルの処理が完了しました:\n"
                    f"元ファイル: {os.path.basename(file_path)}\n"
                    f"新ファイル: {os.path.basename(new_file_path)}\n"
                    f"置換回数: {total_replacements}回\n"
                    f"置換内容: {replacement_text}",
                )
                return True, new_file_path
            else:
                messagebox.showwarning(
                    "警告",
                    f"置換対象のキー文字列が見つかりませんでした。\n"
                    f"検索対象キー文字列: {', '.join(key_strings)}",
                )
                return False, None

        except Exception as e:
            messagebox.showerror("エラー", f"Wordファイルの処理に失敗しました:\n{str(e)}")
            return False, None

    def process_word_direct(self):
        """入力データを直接Wordファイルに適用"""
        try:
            # 入力値の取得
            username = self.username_var.get().strip()
            model = self.model_var.get().strip()
            manufacturing = self.manufacturing_var.get().strip()
            order = self.order_var.get().strip()

            # 入力値の検証
            validation_result = self.validate_inputs(username, model, manufacturing, order)
            if not validation_result["valid"]:
                messagebox.showerror("入力エラー", validation_result["message"])
                return

            # Wordファイルを処理
            success, new_file_path = self.process_word_file(username, model, manufacturing, order)

            if success:
                # 結果表示エリアに更新
                self.display_result(username, model, manufacturing, order)

                # 追加のWord処理完了メッセージ
                word_result = f"""\n\n📝 Word処理完了:
✅ 元ファイル: check2.docx
✅ 新ファイル: {os.path.basename(new_file_path)}
✅ キー文字列「検査対象情報」を「{order}/{manufacturing}」に置換
✅ 新しいファイルとして保存されました"""

                self.result_text.insert(tk.END, word_result)

        except Exception as e:
            messagebox.showerror("エラー", f"Word処理中にエラーが発生しました: {str(e)}")

    # ==================== メイン実行メソッド ====================

    def run(self):
        """GUIの実行"""
        self.root.mainloop()


# ==================== メイン実行部分 ====================

if __name__ == "__main__":
    rpa = NewRPASystem()
    rpa.run()